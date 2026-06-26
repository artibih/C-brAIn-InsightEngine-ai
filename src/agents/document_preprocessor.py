import json
import re
import os
from pathlib import Path
from typing import List, Optional
import subprocess
import tempfile
import zipfile
import shutil
import base64
from urllib.parse import urlparse

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage
from config.llm_config import get_llm
from config.llm_selection import ResolvedLlmSelection
from src.prompts.document_preprocessor_prompt import DOCUMENT_SECTION_PROMPT
import structlog

logger = structlog.get_logger()


class DocumentPreprocessor:
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.tex', '.latex', '.zip'}
    MAX_IMAGES_PER_DOC = 10

    def __init__(
        self,
        llm_selection: ResolvedLlmSelection | None = None,
        azure_connection_string: Optional[str] = None,
    ):
        self.llm = get_llm(
            temperature=0,
            workload="online",
            llm_selection=llm_selection,
        )
        self.prompt = ChatPromptTemplate.from_template(DOCUMENT_SECTION_PROMPT)
        self.temp_dirs = []
        self.temp_files = []
        self.azure_connection_string = azure_connection_string or os.getenv('AZURE_STORAGE_CONNECTION_STRING') 
    
    def __del__(self):
        """Cleanup temporary directories and files on destruction."""
        for temp_dir in self.temp_dirs:
            try:
                if os.path.exists(temp_dir):
                    shutil.rmtree(temp_dir)
            except Exception as e:
                logger.warning("temp_dir_cleanup_failed", temp_dir=temp_dir, error=str(e))

        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                logger.warning("temp_file_cleanup_failed", temp_file=temp_file, error=str(e))


    async def _describe_image_with_llm(self, image_bytes: bytes, mime_type: str = "image/png") -> Optional[str]:
        """Use vision LLM to describe an image from a scientific paper."""
        image_b64 = base64.b64encode(image_bytes).decode("utf-8")

        message = HumanMessage(content=[
            {
                "type": "image_url",
                "image_url": {
                    "url": f"data:{mime_type};base64,{image_b64}"
                }
            },
            {
                "type": "text",
                "text": (
                    "You are analyzing a figure from a scientific paper. "
                    "Describe what this figure shows in detail: any charts, graphs, diagrams, "
                    "tables, or visual data. Include axis labels, trends, key values, and conclusions "
                    "a reviewer would need to evaluate this figure."
                )
            }
        ])

        try:
            response = await self.llm.ainvoke([message])
            return response.content
        except Exception as e:
            logger.error("vision_llm_failed", error=str(e))
            return None


    def _is_azure_blob_url(self, path: str) -> bool:
        return path.startswith('https://') and '.blob.core.windows.net' in path

    def _download_from_azure(self, blob_url: str) -> Optional[str]:
        """Download a file from Azure Blob Storage to a temporary local file."""
        try:
            from azure.storage.blob import BlobServiceClient

            logger.info("downloading_from_azure", blob_url=blob_url)

            parsed = urlparse(blob_url)
            path_parts = parsed.path.lstrip('/').split('/', 1)

            if len(path_parts) < 2:
                logger.error("invalid_blob_url_format", blob_url=blob_url)
                return None

            container_name = path_parts[0]
            blob_name = path_parts[1]
            account_name = parsed.netloc.split('.')[0]

            logger.info(
                "parsed_blob_url",
                account=account_name,
                container=container_name,
                blob=blob_name
            )

            if self.azure_connection_string:
                blob_service_client = BlobServiceClient.from_connection_string(
                    self.azure_connection_string
                )
            else:
                from azure.identity import DefaultAzureCredential
                account_url = f"https://{account_name}.blob.core.windows.net"
                blob_service_client = BlobServiceClient(
                    account_url=account_url,
                    credential=DefaultAzureCredential()
                )

            blob_client = blob_service_client.get_blob_client(
                container=container_name,
                blob=blob_name
            )

            file_ext = Path(blob_name).suffix
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp:
                tmp_path = tmp.name

            with open(tmp_path, 'wb') as f:
                download_stream = blob_client.download_blob()
                f.write(download_stream.readall())

            self.temp_files.append(tmp_path)

            logger.info(
                "azure_download_successful",
                blob_url=blob_url,
                local_path=tmp_path,
                size=os.path.getsize(tmp_path)
            )

            return tmp_path

        except ImportError:
            logger.error("azure_storage_blob_not_installed")
            return None
        except Exception as e:
            logger.error("azure_download_failed", blob_url=blob_url, error=str(e), exc_info=True)
            return None


    async def process_files(self, file_paths: List[str]) -> dict:
        """
        Download (if Azure URLs), read, and process documents into sections.
        Supports PDF, DOCX, DOC, TEX, LATEX, ZIP.
        """
        logger.info("starting_file_processing", file_count=len(file_paths))

        local_paths = []
        for file_path in file_paths:
            if self._is_azure_blob_url(file_path):
                local_path = self._download_from_azure(file_path)
                if local_path:
                    local_paths.append(local_path)
                else:
                    logger.error("failed_to_download_azure_file", blob_url=file_path)
            else:
                local_paths.append(file_path)

        if not local_paths:
            logger.error("no_valid_files_to_process")
            return {"sections": [], "error": "Failed to download or locate any documents"}

        all_texts = []
        for path in local_paths:
            text = await self.read_single_document_async(path)
            if text:
                all_texts.append(text)
                logger.info("document_read_successfully", file_path=path, text_length=len(text))
            else:
                logger.warning("document_read_returned_empty", file_path=path)

        if not all_texts:
            logger.error("all_document_reads_failed", file_count=len(local_paths))
            return {"sections": [], "error": "Failed to read any documents"}

        combined = "\n\n" + "=" * 80 + "\n\n".join(all_texts)
        logger.info("documents_combined", document_count=len(all_texts), total_length=len(combined))

        return await self.process_text(combined)


    async def read_single_document_async(self, file_path: str) -> Optional[str]:
        """Route a local file to the appropriate async reader."""
        if not os.path.exists(file_path):
            logger.error("file_not_found", file_path=file_path)
            return None

        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.warning("file_is_empty", file_path=file_path)
            return None

        ext = Path(file_path).suffix.lower()
        logger.info("reading_file_with_extension", file_path=file_path, extension=ext, size_bytes=file_size)

        if ext not in self.SUPPORTED_EXTENSIONS:
            logger.error(
                "unsupported_file_extension",
                file_path=file_path,
                extension=ext,
                supported=list(self.SUPPORTED_EXTENSIONS)
            )
            return None

        try:
            if ext == '.pdf':
                return await self._read_pdf_async(file_path)
            elif ext in {'.docx', '.doc'}:
                return await self._read_docx_async(file_path)
            elif ext in {'.tex', '.latex'}:
                return await self._read_latex_async(file_path, project_root=None)
            elif ext == '.zip':
                return await self._read_zip_async(file_path)
        except Exception as e:
            logger.error(
                "document_reader_exception",
                file_path=file_path,
                extension=ext,
                error=str(e),
                exc_info=True
            )
            return None

   
    async def _read_pdf_async(self, file_path: str) -> Optional[str]:
        """Extract text and describe images from a PDF using pymupdf + vision LLM."""
        try:
            import fitz  

            all_parts = []
            image_count = 0

            with fitz.open(file_path) as doc:
                num_pages = len(doc)
                logger.info("reading_pdf", file_path=file_path, num_pages=num_pages)

                for page_num, page in enumerate(doc):

                    text = page.get_text()
                    if text.strip():
                        all_parts.append(f"\n--- Page {page_num + 1} ---\n{text}")

                    if image_count >= self.MAX_IMAGES_PER_DOC:
                        logger.warning("pdf_image_limit_reached", limit=self.MAX_IMAGES_PER_DOC)
                        break

                    for img in page.get_images(full=True):
                        if image_count >= self.MAX_IMAGES_PER_DOC:
                            break

                        xref = img[0]
                        try:
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            ext = base_image.get("ext", "png")
                            mime = f"image/{ext}"

                            description = await self._describe_image_with_llm(image_bytes, mime)
                            if description:
                                all_parts.append(
                                    f"\n[FIGURE on page {page_num + 1}]\n{description}\n[/FIGURE]\n"
                                )
                                image_count += 1
                        except Exception as e:
                            logger.warning("pdf_image_extraction_failed", xref=xref, error=str(e))
                            continue

            return "\n".join(all_parts) if all_parts else None

        except ImportError:
            logger.warning("pymupdf_not_installed_falling_back")
            return self._read_pdf_fallback(file_path)
        except Exception as e:
            logger.error("pdf_read_failed", file_path=file_path, error=str(e))
            return self._read_pdf_fallback(file_path)

    def _read_pdf_fallback(self, file_path: str) -> Optional[str]:
        """Fallback PDF text extraction using pypdf, then pdftotext CLI."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            all_text = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        all_text.append(f"\n--- Page {i + 1} ---\n{text}")
                except Exception as e:
                    logger.warning("pypdf_page_failed", page=i + 1, error=str(e))
                    continue

            if all_text:
                logger.info("pdf_read_via_pypdf", file_path=file_path)
                return "\n".join(all_text)
        except Exception:
            pass

        try:
            result = subprocess.run(
                ['pdftotext', file_path, '-'],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                logger.info("pdf_read_via_pdftotext", file_path=file_path)
                return result.stdout
        except Exception as e:
            logger.error("pdf_fallback_failed", file_path=file_path, error=str(e))

        return None



    async def _read_docx_async(self, file_path: str) -> Optional[str]:
        """Extract text and describe images from a DOCX file using vision LLM."""
        ext = Path(file_path).suffix.lower()

        if ext == '.doc':
            return await self._convert_and_read_doc_async(file_path)

        try:
            from docx import Document

            doc = Document(file_path)
            all_parts = []
            image_count = 0

            for para in doc.paragraphs:
                if para.text.strip():
                    all_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        all_parts.append(row_text)

            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    if image_count >= self.MAX_IMAGES_PER_DOC:
                        logger.warning("docx_image_limit_reached", limit=self.MAX_IMAGES_PER_DOC)
                        break
                    try:
                        image_bytes = rel.target_part.blob
                        content_type = rel.target_part.content_type 
                        mime = content_type if content_type.startswith("image/") else "image/png"

                        description = await self._describe_image_with_llm(image_bytes, mime)
                        if description:
                            all_parts.append(f"\n[FIGURE]\n{description}\n[/FIGURE]\n")
                            image_count += 1
                    except Exception as e:
                        logger.warning("docx_image_extraction_failed", error=str(e))
                        continue

            logger.info("docx_read_successfully", file_path=file_path, images_described=image_count)
            return "\n\n".join(all_parts) if all_parts else None

        except Exception as e:
            logger.error("docx_read_failed", file_path=file_path, error=str(e))
            return None

    async def _convert_and_read_doc_async(self, file_path: str) -> Optional[str]:
        """Convert legacy .doc to .docx via LibreOffice, then read."""
        try:
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name

            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'docx',
                 '--outdir', os.path.dirname(tmp_path), file_path],
                capture_output=True, text=True, timeout=120
            )

            if result.returncode == 0:
                text = await self._read_docx_async(tmp_path)
                os.unlink(tmp_path)
                logger.info("doc_converted_and_read", file_path=file_path)
                return text
            else:
                logger.error("doc_conversion_failed", file_path=file_path, stderr=result.stderr)
                return None

        except Exception as e:
            logger.error("doc_conversion_exception", file_path=file_path, error=str(e))
            return None

    async def _read_latex_async(self, file_path: str, project_root: Optional[str] = None) -> Optional[str]:
        """Read a .tex file, describe referenced images with vision LLM, then clean markup."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()

            logger.info("latex_read_successfully", file_path=file_path, length=len(content))

            if project_root:
                content = await self._extract_latex_images_async(content, file_path, project_root)

            return self._clean_latex(content)

        except Exception as e:
            logger.error("latex_read_failed", file_path=file_path, error=str(e))
            return None

    async def _extract_latex_images_async(
        self,
        latex_content: str,
        tex_file_path: str,
        project_root: str
    ) -> str:
        """Find \includegraphics references and describe each image with vision LLM."""
        pattern = r'\\includegraphics(?:\[.*?\])?\{([^}]+)\}'
        matches = list(re.finditer(pattern, latex_content))

        image_descriptions = []

        for match in matches:
            if len(image_descriptions) >= self.MAX_IMAGES_PER_DOC:
                logger.warning("latex_image_limit_reached", limit=self.MAX_IMAGES_PER_DOC)
                break

            image_path = match.group(1)
            resolved_path = self._resolve_latex_image_path(image_path, tex_file_path, project_root)

            if resolved_path and os.path.exists(resolved_path):
                ext = Path(resolved_path).suffix.lower().lstrip(".")
                mime = f"image/{ext}" if ext in {"png", "jpg", "jpeg", "gif", "bmp", "webp"} else "image/png"

                try:
                    with open(resolved_path, "rb") as f:
                        image_bytes = f.read()

                    description = await self._describe_image_with_llm(image_bytes, mime)
                    if description:
                        image_descriptions.append(
                            f"\n[IMAGE: {image_path}]\n{description}\n[/IMAGE]\n"
                        )
                        logger.info("latex_image_described", image_path=image_path)
                except Exception as e:
                    logger.warning("latex_image_read_failed", image_path=image_path, error=str(e))
            else:
                logger.warning("latex_image_not_resolved", image_path=image_path)

        if image_descriptions:
            return (
                latex_content
                + "\n\n" + "=" * 80
                + "\nEXTRACTED IMAGES:\n" + "=" * 80 + "\n"
                + "\n".join(image_descriptions)
            )

        return latex_content

    def _clean_latex(self, content: str) -> str:
        """Strip LaTeX markup, leaving readable prose for the LLM."""
        content = re.sub(r'%.*?\n', '\n', content)
        for env in ['equation', 'align', 'math', 'figure', 'table', 'lstlisting', 'verbatim', 'tikzpicture']:
            content = re.sub(
                rf'\\begin\{{{env}\*?\}}.*?\\end\{{{env}\*?\}}',
                f'[{env.upper()} BLOCK REMOVED]',
                content, flags=re.DOTALL
            )

        content = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}', r'\1', content)
        content = re.sub(r'\\[a-zA-Z]+\*?', '', content)
        content = re.sub(r'[{}]', '', content)
        content = re.sub(r'\n{3,}', '\n\n', content)
        return content.strip()

    def _resolve_latex_image_path(
        self,
        image_path: str,
        tex_file_path: str,
        project_root: str
    ) -> Optional[str]:
        """Resolve a LaTeX image reference to an actual file path."""
        extensions = ['', '.png', '.jpg', '.jpeg', '.pdf', '.eps', '.ps',
                      '.PNG', '.JPG', '.JPEG', '.PDF']
        tex_dir = os.path.dirname(tex_file_path)
        search_dirs = [tex_dir, project_root]

        for base_dir in search_dirs:
            for ext in extensions:
                candidate = os.path.join(base_dir, image_path + ext)
                if os.path.exists(candidate):
                    logger.info("latex_image_resolved", image_path=image_path, resolved=candidate)
                    return candidate

        logger.warning("latex_image_not_found", image_path=image_path, searched_dirs=search_dirs)
        return None


    async def _read_zip_async(self, file_path: str) -> Optional[str]:
        """Extract a ZIP (typically a LaTeX project) and read all documents inside."""
        try:
            logger.info("processing_zip_file", file_path=file_path)
            temp_dir = tempfile.mkdtemp(prefix='latex_project_')
            self.temp_dirs.append(temp_dir)

            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            extracted_files = self._find_documents_in_directory(temp_dir)

            if not extracted_files:
                logger.warning("no_documents_found_in_zip", file_path=file_path)
                return None

            logger.info("found_documents_in_zip", file_path=file_path, count=len(extracted_files))

            all_texts = []
            for doc_path in extracted_files:
                try:
                    rel_path = os.path.relpath(doc_path, temp_dir)
                    text = await self._read_single_document_from_zip_async(doc_path, temp_dir)
                    if text:
                        all_texts.append(f"\n{'=' * 80}\n[File: {rel_path}]\n{'=' * 80}\n{text}")
                        logger.info("zip_document_read", file=rel_path)
                except Exception as e:
                    logger.error("zip_document_read_failed", file_path=doc_path, error=str(e))
                    continue

            return "\n\n".join(all_texts) if all_texts else None

        except zipfile.BadZipFile:
            logger.error("invalid_zip_file", file_path=file_path)
            return None
        except Exception as e:
            logger.error("zip_read_exception", file_path=file_path, error=str(e))
            return None

    async def _read_single_document_from_zip_async(
        self, file_path: str, project_root: str
    ) -> Optional[str]:
        """Read a document extracted from a ZIP, passing project_root for image resolution."""
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return await self._read_pdf_async(file_path)
        elif ext in {'.docx', '.doc'}:
            return await self._read_docx_async(file_path)
        elif ext in {'.tex', '.latex'}:
            return await self._read_latex_async(file_path, project_root=project_root)

        return None

    def _find_documents_in_directory(self, directory: str) -> List[str]:
        """Recursively find all supported documents in a directory."""
        document_files = []

        for root, dirs, files in os.walk(directory):
            dirs[:] = [d for d in dirs if d not in {
                '__MACOSX', '.git', 'node_modules', '.venv', 'venv', '.aux', 'build'
            }]

            for file in files:
                file_path = os.path.join(root, file)
                ext = Path(file_path).suffix.lower()
                if ext in {'.pdf', '.docx', '.doc', '.tex', '.latex'}:
                    document_files.append(file_path)

        document_files.sort(key=lambda x: (
            0 if Path(x).stem in {'main', 'paper', 'manuscript', 'thesis', 'document'} else 1,
            x
        ))

        return document_files

    async def process_text(self, text: str) -> dict:
        if not text or not text.strip():
            logger.error("process_text_called_with_empty_text")
            return {"sections": [], "error": "No text content to process"}

        logger.info(
            "processing_text_with_llm",
            text_length=len(text),
            text_preview=text[:200]
        )

        chain = self.prompt | self.llm
        response = await chain.ainvoke({"manuscript": text})

        logger.info(
            "llm_response_received",
            response_length=len(response.content) if hasattr(response, 'content') else 0
        )

        raw = response.content

        raw = re.sub(r'^```(?:json)?\s*', '', raw.strip())
        raw = re.sub(r'\s*```$', '', raw.strip())

        raw = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', raw)

        try:
            parsed = json.loads(raw)
            logger.info("json_parse_successful", sections_count=len(parsed.get('sections', [])))
            return parsed
        except Exception as e:
            logger.warning("json_parse_failed", error=str(e))
            try:
                match = re.search(r"\{.*\}", raw, re.DOTALL)
                if match:
                    cleaned = re.sub(r'\\(?!["\\/bfnrtu])', r'\\\\', match.group())
                    parsed = json.loads(cleaned)
                    logger.info("json_extracted_from_response", sections_count=len(parsed.get('sections', [])))
                    return parsed
            except Exception as e2:
                logger.error("fallback_json_parse_failed", error=str(e2))

            return {
                "sections": [],
                "raw_output": response.content,
                "error": "Failed to parse LLM response as JSON"
            }